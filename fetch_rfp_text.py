"""
fetch_rfp_text.py — Fetch SAM.gov solicitation attachments for classified AI contracts.

Two-phase approach to minimise SAM.gov API calls (capped at ~10/day free tier):

  Phase 1 (free, no quota): Hit USASpending award detail endpoint for each AI
    contract to get its solicitation_identifier. Stored in
    data/solicitation_ids.json — run once, re-used forever.

  Phase 2 (burns SAM quota): Broad date-range scans of SAM.gov opportunities
    (no NAICS filter, filtered to DHS). Match results against the target
    solicitation numbers. Each call covers up to 1000 opportunities, so 5 years
    of DHS opps typically needs ~15-30 calls total — done in 2-3 days at 10/day.
    Attachments (PDF/DOCX/XLSX) are free S3 downloads that don't count against
    quota.

Run:
    python3 fetch_rfp_text.py --phase1        # get solicitation IDs (free, run first)
    python3 fetch_rfp_text.py --phase2        # SAM.gov scan + download (burns quota)
    python3 fetch_rfp_text.py --phase2 --dry-run   # show matches without downloading
"""

import argparse
import csv
import hashlib
import io
import json
import os
import re
import time
from datetime import date, datetime, timedelta, timezone
from pathlib import Path

import requests
from dotenv import load_dotenv

load_dotenv()

SAM_API_KEY = os.environ.get("SAM_API_KEY")
OPP_SEARCH  = "https://api.sam.gov/prod/opportunities/v2/search"
AWARD_DETAIL = "https://api.usaspending.gov/api/v2/awards/{}//"

RESULTS_JSON     = Path("web/data/results.json")
SOL_IDS_JSON     = Path("data/solicitation_ids.json")
SCAN_CURSOR_JSON = Path("data/rfp_scan_cursor.json")
PROCESSED_JSON   = Path("data/rfp_processed.json")
TEXT_DIR         = Path("data/rfp_texts")

MAX_WINDOW_DAYS  = 364


# ---------------------------------------------------------------------------
# Text extraction (same as rfp_text_pipeline.py)
# ---------------------------------------------------------------------------

def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join(p.extract_text() or "" for p in reader.pages).strip()
    except Exception as e:
        return f"[pdf extract failed: {e}]"

def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text]
        for t in doc.tables:
            for row in t.rows:
                cells = [c.text for c in row.cells if c.text]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts).strip()
    except Exception as e:
        return f"[docx extract failed: {e}]"

def _extract_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        out = []
        for ws in wb.worksheets:
            out.append(f"### Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(v).strip() for v in row if v is not None]
                if cells:
                    out.append(" | ".join(cells))
        return "\n".join(out).strip()
    except Exception as e:
        return f"[xlsx extract failed: {e}]"

def extract_text(filename: str, content_type: str, data: bytes) -> str:
    f = filename.lower()
    ct = (content_type or "").lower()
    if f.endswith(".pdf") or "pdf" in ct:
        return _extract_pdf(data)
    if f.endswith(".docx") or "wordprocessingml" in ct:
        return _extract_docx(data)
    if f.endswith(".xlsx") or "spreadsheetml" in ct:
        return _extract_xlsx(data)
    return ""


# ---------------------------------------------------------------------------
# Phase 1: get solicitation IDs from USASpending (free)
# ---------------------------------------------------------------------------

def phase1_get_solicitation_ids() -> None:
    if not RESULTS_JSON.exists():
        raise SystemExit(f"{RESULTS_JSON} not found — run classify_ai.py + build_web.py first")

    data = json.loads(RESULTS_JSON.read_text())
    contracts = [c for c in data["contracts"] if c.get("generated_internal_id")]

    existing: dict = json.loads(SOL_IDS_JSON.read_text()) if SOL_IDS_JSON.exists() else {}
    print(f"{len(contracts)} AI contracts  |  {len(existing)} already have solicitation IDs")

    session = requests.Session()
    new = 0
    for c in contracts:
        gid = c["generated_internal_id"]
        award_id = c["award_id"]
        if award_id in existing:
            continue
        try:
            r = session.get(AWARD_DETAIL.format(gid), timeout=20)
            if r.status_code != 200:
                print(f"  {award_id}: HTTP {r.status_code}")
                continue
            sol = r.json().get("latest_transaction_contract_data", {}).get("solicitation_identifier")
            existing[award_id] = {"solicitation_id": sol, "generated_internal_id": gid}
            new += 1
            if sol:
                print(f"  {award_id} → {sol}")
            else:
                print(f"  {award_id} → (no solicitation ID)")
            time.sleep(0.3)
        except Exception as e:
            print(f"  {award_id}: {e}")

    SOL_IDS_JSON.parent.mkdir(parents=True, exist_ok=True)
    SOL_IDS_JSON.write_text(json.dumps(existing, indent=2))
    print(f"\nSaved {len(existing)} entries (+{new} new) → {SOL_IDS_JSON}")

    sol_ids = {v["solicitation_id"] for v in existing.values() if v.get("solicitation_id")}
    print(f"{len(sol_ids)} unique solicitation IDs to search for")


# ---------------------------------------------------------------------------
# Phase 2: SAM.gov scan + download
# ---------------------------------------------------------------------------

def _mmddyyyy(d: date) -> str:
    return d.strftime("%m/%d/%Y")

def _year_chunks(start: date, end: date) -> list[tuple[date, date]]:
    chunks, cur = [], start
    while cur < end:
        nxt = min(cur + timedelta(days=MAX_WINDOW_DAYS), end)
        chunks.append((cur, nxt))
        cur = nxt + timedelta(days=1)
    return chunks


def phase2_scan(dry_run: bool = False, max_calls: int = 10) -> None:
    if not SAM_API_KEY:
        raise SystemExit("SAM_API_KEY not set")
    if not SOL_IDS_JSON.exists():
        raise SystemExit(f"{SOL_IDS_JSON} not found — run --phase1 first")

    sol_data = json.loads(SOL_IDS_JSON.read_text())
    target_sol_ids = {
        v["solicitation_id"]
        for v in sol_data.values()
        if v.get("solicitation_id")
    }
    print(f"Targeting {len(target_sol_ids)} solicitation IDs")

    processed: set = set(json.loads(PROCESSED_JSON.read_text()) if PROCESSED_JSON.exists() else [])
    cursor = json.loads(SCAN_CURSOR_JSON.read_text()) if SCAN_CURSOR_JSON.exists() else None

    # Date range: 5 years back
    end_date   = date.today()
    start_date = date(2021, 10, 1)
    chunks     = _year_chunks(start_date, end_date)

    TEXT_DIR.mkdir(parents=True, exist_ok=True)
    session  = requests.Session()
    calls    = 0
    matched  = 0
    downloaded = 0

    # Resume from cursor if present
    resume_chunk_idx = 0
    resume_offset    = 0
    if cursor:
        print(f"Resuming from cursor: {cursor['chunk_from']} offset {cursor['offset']}")
        for i, (cf, ct) in enumerate(chunks):
            if str(cf) == cursor["chunk_from"]:
                resume_chunk_idx = i
                resume_offset    = cursor["offset"]
                break

    for chunk_idx, (chunk_from, chunk_to) in enumerate(chunks):
        if chunk_idx < resume_chunk_idx:
            print(f"  {chunk_from}→{chunk_to}: skipping (before cursor)")
            continue

        offset = resume_offset if chunk_idx == resume_chunk_idx else 0
        resume_offset = 0  # only apply once

        while calls < max_calls:
            calls += 1
            params = {
                "api_key":    SAM_API_KEY,
                "postedFrom": _mmddyyyy(chunk_from),
                "postedTo":   _mmddyyyy(chunk_to),
                "limit":      1000,
                "offset":     offset,
                "deptname":   "Department of Homeland Security",
            }
            try:
                r = session.get(OPP_SEARCH, params=params, timeout=120)
            except requests.RequestException as e:
                print(f"  request error: {e}")
                break

            if r.status_code == 429:
                print("SAM 429 — quota exhausted. Saving cursor.")
                _save_cursor(chunk_from, chunk_to, offset)
                _save_processed(processed)
                return
            if r.status_code != 200:
                print(f"SAM HTTP {r.status_code}: {r.text[:200]}")
                break

            opps = r.json().get("opportunitiesData") or []
            total = r.json().get("totalRecords", 0)
            if offset == 0:
                print(f"  {chunk_from}→{chunk_to}: {total:,} DHS opps")

            for opp in opps:
                sol_num = opp.get("solicitationNumber", "")
                notice_id = opp.get("noticeId", "")
                if sol_num not in target_sol_ids:
                    continue
                if notice_id in processed:
                    continue
                matched += 1
                print(f"  MATCH: {sol_num} | {opp.get('title','')[:60]}")

                if not dry_run:
                    n = _download_attachments(session, opp, sol_num)
                    downloaded += n
                processed.add(notice_id)

            if len(opps) < 1000:
                break  # drained this chunk
            offset += 1000
            time.sleep(1.0)

        if calls >= max_calls:
            print(f"Call cap reached ({max_calls}). Saving cursor.")
            _save_cursor(chunk_from, chunk_to, offset)
            break
    else:
        # Fully drained — clear cursor
        if SCAN_CURSOR_JSON.exists():
            SCAN_CURSOR_JSON.unlink()
        print("Full scan complete — cursor cleared.")

    _save_processed(processed)
    print(f"\nDone: {calls} API calls, {matched} matches, {downloaded} attachments downloaded")


def _download_attachments(session: requests.Session, opp: dict, sol_num: str) -> int:
    links = opp.get("resourceLinks") or []
    out_dir = TEXT_DIR / re.sub(r"[^\w\-]", "_", sol_num)
    out_dir.mkdir(parents=True, exist_ok=True)
    count = 0

    for i, url in enumerate(links):
        fetch_url = url
        if "sam.gov" in url and "api_key=" not in url:
            fetch_url = f"{url}{'&' if '?' in url else '?'}api_key={SAM_API_KEY}"
        try:
            resp = session.get(fetch_url, timeout=180, stream=True)
            if resp.status_code != 200:
                continue
            data = resp.content
            cd   = resp.headers.get("content-disposition", "")
            filename = ""
            if "filename=" in cd:
                filename = cd.split("filename=", 1)[1].strip().strip('"').strip("'")
            if not filename:
                filename = re.split(r"[?#]", url.rstrip("/").split("/")[-1])[0] or f"attachment_{i}"

            text = extract_text(filename, resp.headers.get("content-type", ""), data)
            if not text:
                continue

            out = {
                "sol_num":  sol_num,
                "notice_id": opp.get("noticeId"),
                "title":    opp.get("title"),
                "filename": filename,
                "sha256":   hashlib.sha256(data).hexdigest(),
                "chars":    len(text),
                "text":     text,
            }
            (out_dir / f"{i}_{filename}.json").write_text(json.dumps(out, indent=2))
            count += 1
            print(f"    {filename}: {len(text):,} chars")
        except Exception as e:
            print(f"    attachment {i} failed: {e}")

    return count


def _save_cursor(chunk_from, chunk_to, offset):
    SCAN_CURSOR_JSON.parent.mkdir(parents=True, exist_ok=True)
    SCAN_CURSOR_JSON.write_text(json.dumps({
        "chunk_from": str(chunk_from),
        "chunk_to":   str(chunk_to),
        "offset":     offset,
        "saved_at":   datetime.now(timezone.utc).isoformat(),
    }))


def _save_processed(processed):
    PROCESSED_JSON.parent.mkdir(parents=True, exist_ok=True)
    PROCESSED_JSON.write_text(json.dumps(sorted(processed)))


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--phase1", action="store_true", help="Get solicitation IDs from USASpending (free)")
    ap.add_argument("--phase2", action="store_true", help="SAM.gov scan + attachment download")
    ap.add_argument("--dry-run", action="store_true", help="Phase 2: show matches without downloading")
    ap.add_argument("--max-calls", type=int, default=9, help="SAM API call cap (default 9, saves 1 for safety)")
    args = ap.parse_args()

    if not args.phase1 and not args.phase2:
        ap.print_help()
        return

    if args.phase1:
        phase1_get_solicitation_ids()
    if args.phase2:
        phase2_scan(dry_run=args.dry_run, max_calls=args.max_calls)


if __name__ == "__main__":
    main()
